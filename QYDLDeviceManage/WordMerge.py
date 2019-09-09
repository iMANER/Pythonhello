#pip install python-docx
import docx

#获得文档
doc1=docx.Document("3.docx")
doc2=docx.Document("4.docx")

str1=[]
#按照段落读取内容
for para in doc1.paragraphs:
    str1.append(para.text)
for para in doc2.paragraphs:
    str1.append(para.text)
for s in str1:
    print(s)


from docx import Document
# 合并文档的列表
files = ['1.docx', '2.docx']
#合并操作
def combine_word_documents(files):
    merged_document = Document()

    for index, file in enumerate(files):
        sub_doc = Document(file)

        # Don't add a page break if you've reached the last file.
        if index < len(files)-1:
           sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

    merged_document.save('merged.docx')

combine_word_documents(files)



#pip install docxcompose
import os
from docx import Document
from docxcompose.composer import Composer
from docx import Document as Document_compose
result=[]
def search(path=".", name=""):
    for item in os.listdir(path):
        item_path = os.path.join(path, item)
        if os.path.isdir(item_path):
            search(item_path, name)
        elif os.path.isfile(item_path):
            if name in item:
                global result
                result.append(item_path)
                print (item_path)

#search(path=r"F:\haiyan", name=".docx")
search(path=r"./DeviceInfo", name=".docx")
print(result)
# 合并文档的列表


files = result
def combine_all_docx(filename_master,files_list):
    number_of_sections=len(files_list)
    master = Document_compose(filename_master)
    composer = Composer(master)
    for i in range(1, number_of_sections):
        doc_temp = Document_compose(files_list[i])
        composer.append(doc_temp)
    composer.save("combined_file.docx")
combine_all_docx(result[0],result)
